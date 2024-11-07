from docx import Document

from utils.str_utils import remove_service_chars


def replace_tokens_in_docx(docx: str, token_pairs: dict):
    doc = Document(docx)
    parags = doc.paragraphs
    tables = doc.tables

    for token in token_pairs.keys():
        value = token_pairs[token]

        for parag in parags:
            clean_parag_text = remove_service_chars(parag.text)
            if token in clean_parag_text:
                parag.paragraph_format
                parag.text = parag.text.replace(token, value)

        for table in tables:
            for cell in table._cells:
                clean_cell_text = remove_service_chars(cell.text)
                if token in clean_cell_text:
                    cell.text = cell.text.replace(token, value)

    return doc


if __name__ == "__main__":
    token_pairs = {
        "{med_oid}": "Тут мог быть oui",
        "{med_name}": "Название организации",
        "{doc_date}": "07.11.2024",
    }
    
    changed_doc = replace_tokens_in_docx("doc_templates/Template_RAMD.docx", token_pairs)
    changed_doc.save("test5.docx")
